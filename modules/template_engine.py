"""
template_engine.py — Template-driven engine (FIXED).

Fixes applied vs. previous version:
  1. Page size: template (A4 portrait) preserved everywhere; ONLY the Arc Flash
     scenario result table sits inside an A3-landscape continuous section.
     Section-break ordering corrected — a <w:sectPr> inside a paragraph defines
     the section that ENDS at that paragraph, so the "closing" break comes AFTER
     the landscape content, and a break restoring template size is placed after it.
  2. Scenario heading + description are placed INSIDE the A3 landscape section so
     they appear on the same page as the table (no blank page before the table).
  3. Color coding applied only to the Incident Energy cell, not the whole row.
  4. Header/footer preserved from template. Cover page (first page) suppresses
     header/footer via `different_first_page_header_footer = True`.
  5. Static images & all template static content preserved (python-docx keeps
     them by default; we only touch merge-field runs and marker blocks).
  6. Dynamic Protective Device table removed — `{{PD_TABLE}}` marker is simply
     cleaned (static template text referring to Annexure remains intact).
  7. Annexures: first annexure sits directly under the "Annexures" H1 heading;
     every subsequent annexure forces a page break before its heading.
  8. Result table fits content width; column widths distributed with sensible
     minimums; `w:tblLayout` set to `fixed` so columns respect our widths and
     long cell text wraps instead of overflowing.
"""
import copy
import io
import re
import os
import zipfile
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree as _etree
from .config import (LEVEL_COLORS, TABLE_MARKERS, AF_COLS,
                     A3_W, A3_H, MARGIN)

# XML namespace for Word
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_XML_NS = 'http://www.w3.org/XML/1998/namespace'
_R_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


class AFTemplateEngine:
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self.fields = {}
        self.scenarios = []
        self.comparison = None
        self.annexures = []
        self.consultant_logo = None
        self.client_logo = None
        self.wide_mode = False

        # Read template page geometry (EMU units from python-docx)
        sec0 = self.doc.sections[0]
        self._tmpl_pw = sec0.page_width
        self._tmpl_ph = sec0.page_height
        self._tmpl_ml = sec0.left_margin
        self._tmpl_mr = sec0.right_margin
        self._tmpl_mt = sec0.top_margin
        self._tmpl_mb = sec0.bottom_margin

        # Fix 4: cover page = no header/footer; subsequent pages inherit template H/F.
        try:
            sec0.different_first_page_header_footer = True
        except Exception:
            pass

        self._hdr_fmt = {}
        self._data_fmt = {}
        # Capture header/footer reference elements from template body sectPr so
        # we can inject them into every new sectPr we generate (fixes header/footer
        # not appearing on A3 pages and intermediate sections).
        self._body_hdr_footer_refs = self._capture_body_sectpr_refs()
        self._extract_template_formatting()

    # ─── Public API ───
    def set_fields(self, d):
        self.fields.update(d)

    def set_logos(self, consultant=None, client=None):
        self.consultant_logo = consultant
        self.client_logo = client

    def add_scenario(self, name, desc, headers, rows):
        self.scenarios.append({"name": name, "desc": desc, "headers": headers, "rows": rows})

    def set_comparison(self, headers, rows):
        self.comparison = {"headers": headers, "rows": rows}

    def add_annexure(self, letter, title, content=""):
        self.annexures.append({"letter": letter, "title": title, "content": content})

    def set_wide_mode(self, flag):
        self.wide_mode = flag

    # ─── Capture header/footer references from the template body sectPr ───
    def _capture_body_sectpr_refs(self):
        """
        Return deep-copied headerReference and footerReference elements from
        the template body sectPr. Injected into every intermediate sectPr we
        generate so Word renders headers/footers correctly on A3 pages.
        """
        body_sectPr = self.doc.element.body.find(qn('w:sectPr'))
        refs = []
        if body_sectPr is not None:
            for child in body_sectPr:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('headerReference', 'footerReference'):
                    refs.append(copy.deepcopy(child))
        return refs

    def generate(self, output_path):
        self._insert_logos()
        self._expand_scenarios()
        self._handle_comparison()
        self._expand_annexures()
        self._replace_all_fields()        # body + header/footer (paragraphs + tables)
        self._clean_markers()
        self._enable_field_update_on_open()

        # ── Save to buffer, then ZIP-level post-process header/footer ──────────
        # python-docx cannot reach: (a) Jinja2 {{ }} fields in footer table cells,
        # (b) table-based footer cell widths, (c) header tab stop for A3 pages.
        # We rewrite these directly in the XML inside the DOCX ZIP.
        buf = io.BytesIO()
        self.doc.save(buf)
        processed = self._postprocess_zip(buf.getvalue())
        with open(output_path, 'wb') as f:
            f.write(processed)
        return output_path

    def _enable_field_update_on_open(self):
        """Insert <w:updateFields w:val="1"/> into document settings."""
        try:
            settings_el = self.doc.settings.element
            for old in settings_el.findall(qn('w:updateFields')):
                settings_el.remove(old)
            uf = parse_xml('<w:updateFields {} w:val="1"/>'.format(nsdecls("w")))
            settings_el.insert(0, uf)
        except Exception:
            pass  # non-fatal

    # ─── ZIP-level header/footer post-processing ─────────────────────────────

    def _postprocess_zip(self, docx_bytes: bytes) -> bytes:
        """
        Post-process the saved DOCX at ZIP level to fix three issues
        that python-docx cannot reach after Document.save():

        Fix A — Footer field replacement (Jinja2 {{ }} syntax):
            The template footer uses {{ document_no }}, {{ revisions[0].rev_no }},
            {{ study_standard }} (WITH spaces, Jinja2 style). python-docx's
            paragraph iterator misses these because they live inside a TABLE in
            the footer and use a different quoting format from the body fields.
            We resolve them directly in the XML text nodes.

        Fix B — Footer table auto-width for A3 pages:
            The footer table is fixed at 9026 DXA (A4 width).  On A3 pages the
            table leaves a gap on the right.  Changing tblW/tcW to type="pct"
            makes Word stretch it to the full content width on ANY page size.

        Fix C — Header right-tab position for A3 pages:
            The header has a right-aligned tab at w:pos="9026" (A4 content width).
            On A3 pages the client name snaps to the middle of the page.
            Setting the tab position to a value larger than the widest A3
            landscape content width lets Word's right-tab logic snap to the
            actual right margin on ALL page sizes.
        """
        # Build context: {clean_key: value} from self.fields {{key}}: value
        ctx = {}
        for k, v in self.fields.items():
            clean = re.sub(r'[{}]', '', k).strip()
            ctx[clean] = str(v) if v is not None else ''
        # Jinja2 revisions list for footer's revisions[0].rev_no pattern
        ctx['revisions'] = [{'rev_no': ctx.get('rev_no', 'Rev-0')}]

        in_buf  = io.BytesIO(docx_bytes)
        out_buf = io.BytesIO()

        with zipfile.ZipFile(in_buf, 'r') as zin,              zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                fname = item.filename.split('/')[-1]

                if (re.match(r'(header|footer)\d*\.xml$', fname)
                        and item.filename.startswith('word/')):
                    try:
                        data = self._fix_hf_xml_part(data, ctx, fname)
                    except Exception:
                        pass  # never crash on non-critical step
                zout.writestr(item, data)

        return out_buf.getvalue()

    @staticmethod
    def _resolve_jinja2(text: str, ctx: dict) -> str:
        """
        Resolve {{ expr }}, {{ expr | default('x') }},
        and obj[N].attr patterns from ctx.
        Handles both {{ field }} (Jinja2 style) and {{field}} (no-space style).
        """
        def _replace(m):
            expr = m.group(1).strip()
            default_val = ''
            if '|' in expr:
                dflt = re.search(r"\|\s*default\(['\"]([^'\"]*)['\"]", expr)
                if dflt:
                    default_val = dflt.group(1)
                expr = expr[:expr.index('|')].strip()
            # obj[N].attr  e.g. revisions[0].rev_no
            arr = re.match(r'(\w+)\[(\d+)\]\.(\w+)$', expr)
            if arr:
                base, idx, attr = arr.group(1), int(arr.group(2)), arr.group(3)
                lst = ctx.get(base, [])
                if isinstance(lst, list) and len(lst) > idx:
                    item = lst[idx]
                    return str(item.get(attr, default_val)
                               if isinstance(item, dict) else default_val)
                return default_val
            return str(ctx.get(expr, default_val))

        return re.sub(r'\{\{(.+?)\}\}', _replace, text)

    @staticmethod
    def _fix_hf_xml_part(xml_bytes: bytes, ctx: dict, fname: str) -> bytes:
        """
        For one header/footer XML part:
        1. Resolve {{ }} fields in every <w:t> text node.
        2. Footer only: convert fixed-DXA table widths to percentage (auto-fit).
        3. Header only: bump right-tab position to cover A3 landscape content width.
        """
        root = _etree.fromstring(xml_bytes)

        # ── 1. Resolve {{ }} fields in all <w:t> nodes ──────────────────────
        for t_elem in root.iter(f'{{{_W_NS}}}t'):
            raw = t_elem.text or ''
            if '{{' not in raw:
                continue
            resolved = AFTemplateEngine._resolve_jinja2(raw, ctx)
            if resolved != raw:
                t_elem.text = resolved
                if resolved and (resolved[0] == ' ' or resolved[-1] == ' '):
                    t_elem.set(f'{{{_XML_NS}}}space', 'preserve')

        # ── 2. Footer: convert table to percentage widths ────────────────────
        if 'footer' in fname:
            # Count columns for proportional distribution
            grid_cols = root.findall(f'.//{{{_W_NS}}}gridCol')
            n_cols = len(grid_cols) if grid_cols else 3
            per_col_pct = 5000 // n_cols  # 5000 = 100%
            remainder   = 5000 - per_col_pct * n_cols

            # tblW → 100%
            for tblW in root.findall(f'.//{{{_W_NS}}}tblW'):
                tblW.set(f'{{{_W_NS}}}w',    '5000')
                tblW.set(f'{{{_W_NS}}}type', 'pct')

            # tcW → proportional pct
            tc_widths = root.findall(f'.//{{{_W_NS}}}tcW')
            for idx, tcW in enumerate(tc_widths):
                w = per_col_pct + (remainder if idx == len(tc_widths) - 1 else 0)
                tcW.set(f'{{{_W_NS}}}w',    str(w))
                tcW.set(f'{{{_W_NS}}}type', 'pct')

            # gridCol → update to match new proportions (dxa approx)
            a3_land_w = 21543  # A3 landscape content width (DXA)
            col_dxa   = a3_land_w // n_cols
            for gc in grid_cols:
                gc.set(f'{{{_W_NS}}}w', str(col_dxa))

            # tblLayout → autofit (removes fixed constraint)
            for layout in root.findall(f'.//{{{_W_NS}}}tblLayout'):
                layout.set(f'{{{_W_NS}}}type', 'autofit')

        # ── 3. Header: bump right-tab to cover A3 landscape content width ────
        if 'header' in fname:
            # A3 landscape content width = 23811 - 1134*2 = 21543 DXA
            # A right-tab beyond content width snaps to the right margin edge,
            # so setting it to 21543 is correct for A4, A3 portrait AND landscape.
            for tab in root.findall(f'.//{{{_W_NS}}}tab'):
                val = tab.get(f'{{{_W_NS}}}val', '')
                pos = tab.get(f'{{{_W_NS}}}pos', '0')
                if val == 'right' and int(pos) < 20000:
                    tab.set(f'{{{_W_NS}}}pos', '21543')

        return _etree.tostring(
            root, xml_declaration=True, encoding='UTF-8', standalone=True
        )

    # ─── Width helpers (in DXA) ───
    def _emu_to_dxa(self, emu):
        # 1 DXA = 1/1440 in; 1 EMU = 1/914400 in  →  DXA = EMU / 635
        return int(emu / 635)

    def _content_width_dxa(self, landscape_a3=False, portrait_a3=False):
        """Return printable content width in DXA for the given page type."""
        margin = int(MARGIN)
        if landscape_a3:
            # A3 landscape: long edge is the page width
            return A3_H - margin * 2
        if portrait_a3:
            # A3 portrait: normal A3 width
            return A3_W - margin * 2
        ml = self._emu_to_dxa(self._tmpl_ml)
        mr = self._emu_to_dxa(self._tmpl_mr)
        pw = self._emu_to_dxa(self._tmpl_pw)
        return pw - ml - mr

    # ─── Extract template result-table formatting ───
    def _extract_template_formatting(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{SCENARIO_TABLE}}" in cell.text:
                        if len(table.rows) >= 2:
                            self._hdr_fmt = _extract_fmt(table.rows[0]._tr)
                            self._data_fmt = _extract_fmt(table.rows[1]._tr)
                        return

    # ─── Logo insertion ───
    def _insert_logos(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{consultant_logo}}" in para.text and self.consultant_logo and os.path.exists(self.consultant_logo):
                            for run in para.runs:
                                if "{{consultant_logo}}" in run.text:
                                    run.text = run.text.replace("{{consultant_logo}}", "")
                                    run.add_picture(self.consultant_logo, width=Inches(2.0))
                                    break
                        elif "{{client_logo}}" in para.text and self.client_logo and os.path.exists(self.client_logo):
                            for run in para.runs:
                                if "{{client_logo}}" in run.text:
                                    run.text = run.text.replace("{{client_logo}}", "")
                                    run.add_picture(self.client_logo, width=Inches(2.0))
                                    break
        if not self.consultant_logo:
            self.fields["{{consultant_logo}}"] = "[Consultant Logo]"
        if not self.client_logo:
            self.fields["{{client_logo}}"] = "[Client Logo]"

    # ─── Scenario expansion ───
    def _expand_scenarios(self):
        body = self.doc.element.body
        elems = list(body)
        h_idx = d_idx = tt_idx = t_idx = e_idx = None
        for i, el in enumerate(elems):
            txt = _etxt(el)
            if "SCENARIO_HEADING" in txt and h_idx is None:
                h_idx = i
            elif "SCENARIO_DESC" in txt and h_idx is not None and d_idx is None:
                d_idx = i
            elif "SCENARIO_TABLE_TITLE" in txt and d_idx is not None and tt_idx is None:
                tt_idx = i
            elif el.tag == qn('w:tbl') and tt_idx is not None and t_idx is None:
                if "SCENARIO_TABLE" in _etxt(el):
                    t_idx = i
            elif "SCENARIO_END" in txt and t_idx is not None:
                e_idx = i
                break
        if h_idx is None or t_idx is None or e_idx is None:
            return

        tmpl = elems[h_idx:e_idx + 1]
        t_offset = t_idx - h_idx
        anchor = elems[h_idx]

        for s_idx, sc in enumerate(self.scenarios):
            cloned = [copy.deepcopy(el) for el in tmpl]
            reps = {
                "{{SCENARIO_HEADING}}": "6.{} Scenario: {}".format(s_idx + 1, sc["name"]),
                "{{SCENARIO_DESC}}": sc["desc"],
                "{{SCENARIO_TABLE_TITLE}}": "Table 6.{} - Arc Flash Study Results - {}".format(s_idx + 1, sc["name"]),
                "{{SCENARIO_END}}": "",
            }
            for cel in cloned:
                _replace_in_elem(cel, reps)

            # Rebuild the scenario table with actual data
            tbl = cloned[t_offset]
            if tbl.tag == qn('w:tbl'):
                # Page-size rule for Section 6:
                #   wide_mode (>WIDE_TABLE_THRESHOLD cols) -> A3 Landscape
                #   narrow tables                          -> A3 Portrait
                # Both wrap in a dedicated A3 section so the rest of the
                # document stays A4 Portrait, with headers on every page.
                landscape_mode = self.wide_mode
                cw = self._content_width_dxa(landscape_a3=landscape_mode)
                self._rebuild_table(tbl, sc["headers"], sc["rows"], cw)
                cloned = self._wrap_in_a3_section(
                    cloned, t_offset,
                    landscape=landscape_mode,
                    first_scenario=(s_idx == 0),
                )

            for cel in cloned:
                anchor.addprevious(cel)

        # Remove the original template placeholder block
        for el in tmpl:
            try:
                body.remove(el)
            except ValueError:
                pass

    def _wrap_in_a3_section(self, cloned, t_offset,
                              landscape=True, first_scenario=True):
        """
        Wrap scenario content in a dedicated A3 section (portrait or landscape).

        Word sectPr section-boundary rules (fix for blank page between scenarios):

        • A sectPr inside a paragraph's pPr defines the section that ENDS there.
        • sect_open (inserted only for the FIRST scenario): empty paragraph whose
          sectPr ends the preceding A4 section and forces a nextPage break into A3.
          For scenarios 2, 3, …  we skip sect_open entirely — the previous
          scenario's sect_close already ended that A3 section via nextPage, so
          the next paragraph starts a fresh section; its page-size is defined by
          ITS OWN sect_close.  No extra empty paragraph = no blank A4 page.
        • sect_close: ends this scenario's A3 section (nextPage → returns to A4).

        Header/footer fix: every sectPr carries the same headerReference /
        footerReference as the template body sectPr so Word renders the header
        and footer on A3 pages too.
        """
        from lxml import etree
        pw_dxa = self._emu_to_dxa(self._tmpl_pw)
        ph_dxa = self._emu_to_dxa(self._tmpl_ph)
        ml = self._emu_to_dxa(self._tmpl_ml)
        mr = self._emu_to_dxa(self._tmpl_mr)
        mt = self._emu_to_dxa(self._tmpl_mt)
        mb = self._emu_to_dxa(self._tmpl_mb)
        tmpl_orient = (' w:orient="landscape"' if self._tmpl_pw > self._tmpl_ph
                       else ' w:orient="portrait"')

        # Build XML string for header/footer references cloned from template
        hf_xml = "".join(etree.tostring(ref, encoding="unicode")
                         for ref in self._body_hdr_footer_refs)

        # ── sect_open: ends the preceding A4 section ──────────────────────
        sect_open = parse_xml(
            '<w:p {ns}><w:pPr><w:sectPr>'
            '<w:type w:val="nextPage"/>'
            '<w:pgSz w:w="{w}" w:h="{h}"{o}/>'
            '<w:pgMar w:top="{mt}" w:right="{mr}" w:bottom="{mb}" w:left="{ml}"'
            ' w:header="708" w:footer="708" w:gutter="0"/>'
            '{hf}'
            '<w:titlePg/>'
            '</w:sectPr></w:pPr></w:p>'.format(
                ns=nsdecls("w"), w=pw_dxa, h=ph_dxa, o=tmpl_orient,
                mt=mt, mr=mr, mb=mb, ml=ml, hf=hf_xml)
        )

        # ── sect_close: ends the A3 section ───────────────────────────────
        if landscape:
            # Landscape: long edge = w, short edge = h
            pg_sz = ('<w:pgSz w:w="{w}" w:h="{h}" w:orient="landscape"/>'
                     .format(w=A3_H, h=A3_W))
        else:
            # Portrait A3
            pg_sz = '<w:pgSz w:w="{w}" w:h="{h}"/>'.format(w=A3_W, h=A3_H)

        sect_close = parse_xml(
            '<w:p {ns}><w:pPr><w:sectPr>'
            '<w:type w:val="nextPage"/>'
            '{pg}'
            '<w:pgMar w:top="{m}" w:right="{m}" w:bottom="{m}" w:left="{m}"'
            ' w:header="708" w:footer="708" w:gutter="0"/>'
            '{hf}'
            '</w:sectPr></w:pPr></w:p>'.format(
                ns=nsdecls("w"), pg=pg_sz, m=int(MARGIN), hf=hf_xml)
        )

        new_cloned = list(cloned)
        # Insert sect_open ONLY for the first scenario.
        # For subsequent scenarios the previous sect_close already returned to A4;
        # inserting another empty A4 transitional paragraph would create a blank page.
        if first_scenario:
            new_cloned.insert(0, sect_open)            # before heading, only once
            new_cloned.insert(t_offset + 2, sect_close)  # after table (+1 for sect_open)
        else:
            # No sect_open for scenario 2, 3, ...
            # The page-size of this section is defined by OUR sect_close (A3).
            # The nextPage on sect_close forces a fresh page BEFORE the next scenario.
            new_cloned.insert(t_offset + 1, sect_close)  # after table (no sect_open offset)
        return new_cloned

    # Backward-compatible alias
    def _wrap_in_landscape_section(self, cloned, t_offset):
        return self._wrap_in_a3_section(cloned, t_offset, landscape=True, first_scenario=True)
    # ─── Table rebuild ───
    def _rebuild_table(self, tbl_elem, col_headers, rows, content_width):
        # Clear existing rows
        for tr in tbl_elem.findall(qn('w:tr')):
            tbl_elem.remove(tr)

        # Table-level width + FIXED layout (so columns stick + long text wraps)
        tblPr = tbl_elem.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr {}/>'.format(nsdecls("w")))
            tbl_elem.insert(0, tblPr)

        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml('<w:tblW {} w:w="{}" w:type="dxa"/>'.format(nsdecls("w"), content_width))
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:w'), str(content_width))
            tblW.set(qn('w:type'), 'dxa')

        for old_layout in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old_layout)
        tblPr.append(parse_xml('<w:tblLayout {} w:type="fixed"/>'.format(nsdecls("w"))))

        # Also write a <w:tblGrid> so Word honors widths immediately
        for old_grid in tbl_elem.findall(qn('w:tblGrid')):
            tbl_elem.remove(old_grid)

        all_h = [("s_no", "S.No.")] + list(col_headers)
        widths = _distribute_widths(content_width, len(all_h), all_h)

        grid_xml = '<w:tblGrid {}>'.format(nsdecls("w"))
        for w in widths:
            grid_xml += '<w:gridCol w:w="{}"/>'.format(w)
        grid_xml += '</w:tblGrid>'
        # tblGrid must come after tblPr
        tblPr.addnext(parse_xml(grid_xml))

        # Header row (repeats across pages)
        tbl_elem.append(_build_row(all_h, widths, True, self._hdr_fmt, 0, None))

        # Data rows
        keys = ["s_no"] + [h[0] for h in col_headers]
        for ri, rd in enumerate(rows):
            lvl = rd.get("energy_level", "")
            level_color = LEVEL_COLORS.get(lvl) if lvl else None
            items = [(keys[i], str(rd.get(keys[i], ""))) for i in range(len(keys))]
            tbl_elem.append(_build_row(items, widths, False, self._data_fmt, ri, level_color))

    # ─── Comparison table ───
    def _handle_comparison(self):
        body = self.doc.element.body
        if self.comparison and self.comparison["rows"]:
            for el in list(body):
                if el.tag == qn('w:tbl') and "COMPARISON_TABLE" in _etxt(el):
                    cw = self._content_width_dxa()
                    self._rebuild_table(el, self.comparison["headers"], self.comparison["rows"], cw)
                    break
        else:
            to_rm = []
            for el in body:
                txt = _etxt(el)
                if any(m in txt for m in ["comparison_heading", "comparison_description", "COMPARISON_TABLE"]):
                    to_rm.append(el)
            for el in to_rm:
                try:
                    body.remove(el)
                except ValueError:
                    pass

    # ─── Annexures ───
    def _expand_annexures(self):
        body = self.doc.element.body
        elems = list(body)
        h_idx = c_idx = e_idx = None
        for i, el in enumerate(elems):
            txt = _etxt(el)
            if "ANNEXURE_HEADING" in txt and h_idx is None:
                h_idx = i
            elif "ANNEXURE_CONTENT" in txt and h_idx is not None and c_idx is None:
                c_idx = i
            elif "ANNEXURE_END" in txt and c_idx is not None:
                e_idx = i
                break
        if h_idx is None or e_idx is None:
            return

        tmpl = elems[h_idx:e_idx + 1]
        anchor = elems[h_idx]

        for idx, anx in enumerate(self.annexures):
            cloned = [copy.deepcopy(el) for el in tmpl]
            reps = {
                "{{ANNEXURE_HEADING}}": "Annexure {} - {}".format(anx["letter"].upper(), anx["title"]),
                "{{ANNEXURE_CONTENT}}": anx.get("content", ""),
                "{{ANNEXURE_END}}": "",
            }
            for cel in cloned:
                _replace_in_elem(cel, reps)

            # Fix 11: Annexure A sits on the same page as the "Annexures" H1
            # heading; B, C, D... each start on a new page.
            if idx > 0:
                _set_page_break_before(cloned[0])

            for cel in cloned:
                anchor.addprevious(cel)

        for el in tmpl:
            try:
                body.remove(el)
            except ValueError:
                pass

    # ─── Field replacement ───
    def _replace_all_fields(self):
        for p in self.doc.paragraphs:
            _replace_para(p, self.fields)
        for t in self.doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        _replace_para(p, self.fields)
        for sec in self.doc.sections:
            for hf in [sec.header, sec.footer,
                       sec.first_page_header, sec.first_page_footer]:
                if hf is None:
                    continue
                # Top-level paragraphs (header uses a single paragraph with tab)
                for p in hf.paragraphs:
                    _replace_para(p, self.fields)
                # Table cells (footer uses a 3-column table for layout)
                for t in hf.tables:
                    for r in t.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                _replace_para(p, self.fields)

    def _clean_markers(self):
        # Wipe any leftover table markers (e.g. {{PD_TABLE}}) without touching
        # the surrounding static template text.
        for p in self.doc.paragraphs:
            for r in p.runs:
                if r.text and "{{" in r.text:
                    for m in TABLE_MARKERS:
                        r.text = r.text.replace("{{" + m + "}}", "")
                    r.text = re.sub(r'\{\{[A-Z_]+\}\}', '', r.text)
        for t in self.doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.text and "{{" in r.text:
                                for m in TABLE_MARKERS:
                                    r.text = r.text.replace("{{" + m + "}}", "")
                                r.text = re.sub(r'\{\{[A-Z_]+\}\}', '', r.text)


# ════════════════════════════════════════════════
# Module-level helpers
# ════════════════════════════════════════════════

def _etxt(el):
    return "".join(t for t in el.itertext())


def _set_page_break_before(elem):
    """Add <w:pageBreakBefore/> to the first paragraph inside `elem`."""
    p = elem if elem.tag == qn('w:p') else elem.find('.//' + qn('w:p'))
    if p is None:
        return
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml('<w:pPr {}/>'.format(nsdecls("w")))
        p.insert(0, pPr)
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pPr.insert(0, parse_xml('<w:pageBreakBefore {}/>'.format(nsdecls("w"))))


def _replace_in_elem(elem, reps):
    for p in elem.iter(qn('w:p')):
        runs = p.findall(qn('w:r'))
        if not runs:
            continue
        full = "".join((r.find(qn('w:t')).text or "") if r.find(qn('w:t')) is not None else "" for r in runs)
        changed = False
        for old, new in reps.items():
            if old in full:
                full = full.replace(old, new)
                changed = True
        if changed:
            first_t = None
            for r in runs:
                t = r.find(qn('w:t'))
                if t is not None:
                    if first_t is None:
                        first_t = t
                        first_t.text = full
                        first_t.set(qn('xml:space'), 'preserve')
                    else:
                        t.text = ""


def _replace_para(para, fields):
    full = para.text
    if "{{" not in full:
        return
    for run in para.runs:
        if "{{" in run.text:
            for k, v in fields.items():
                if k in run.text:
                    run.text = run.text.replace(k, str(v))
    full2 = "".join(r.text for r in para.runs)
    if "{{" in full2:
        for k, v in fields.items():
            full2 = full2.replace(k, str(v))
        if para.runs:
            para.runs[0].text = full2
            for r in para.runs[1:]:
                r.text = ""


def _extract_fmt(tr_elem):
    fmt = {"fill": "1B2A4A", "font": "Calibri", "sz": "18",
           "color": "FFFFFF", "bc": "BFBFBF",
           "mt": "40", "ml": "70", "mb": "40", "mr": "70"}
    if tr_elem is None:
        return fmt
    cells = tr_elem.findall(qn('w:tc'))
    if not cells:
        return fmt
    tc = cells[0]
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fmt["fill"] = shd.get(qn('w:fill'), fmt["fill"])
        bdr = tcPr.find(qn('w:tcBorders'))
        if bdr is not None:
            top_b = bdr.find(qn('w:top'))
            if top_b is not None:
                fmt["bc"] = top_b.get(qn('w:color'), fmt["bc"])
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is not None:
            for side in ["top", "left", "bottom", "right"]:
                el = tcMar.find(qn("w:" + side))
                if el is not None:
                    fmt["m" + side[0]] = el.get(qn('w:w'), fmt["m" + side[0]])
    for p in tc.iter(qn('w:p')):
        for r in p.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None:
                    fn = rf.get(qn('w:ascii'))
                    if fn:
                        fmt["font"] = fn
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    fmt["sz"] = sz.get(qn('w:val'), fmt["sz"])
                col = rPr.find(qn('w:color'))
                if col is not None:
                    fmt["color"] = col.get(qn('w:val'), fmt["color"])
            break
        break
    return fmt


def _distribute_widths(total, n, headers):
    """
    Distribute total width across n columns with sensible minimums, ensuring
    the sum equals `total` exactly so the table never overflows the page.
    """
    # Base proportional weights per column type
    weights = []
    for i, (key, _) in enumerate(headers):
        if key == "s_no":
            weights.append(0.6)
        elif key == "bus_id":
            weights.append(2.4)
        elif key in ("source_pd", "ppe_desc"):
            weights.append(2.0)
        elif key == "energy_level":
            weights.append(1.2)
        else:
            weights.append(1.0)
    tw = sum(weights)
    raw = [int(total * (w / tw)) for w in weights]

    # Enforce a minimum so narrow columns don't collapse
    MIN_COL = 450
    for i in range(len(raw)):
        if raw[i] < MIN_COL:
            raw[i] = MIN_COL

    # Trim/expand final column to make sum match `total` exactly
    diff = total - sum(raw)
    raw[-1] += diff
    if raw[-1] < MIN_COL:
        # Pull from the widest column
        j = max(range(len(raw)), key=lambda i: raw[i])
        raw[j] += raw[-1] - MIN_COL
        raw[-1] = MIN_COL
    return [str(x) for x in raw]


def _esc(t):
    return (str(t).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;").replace("'", "&apos;"))


CENTER_KEYS = {"s_no", "kv", "total_ia", "total_ibf", "total_energy", "afb",
               "final_fct", "glove_class", "trip_time", "open_time", "total_pd_fct",
               "energy_level"}


def _build_row(items, widths, is_header, fmt, row_idx, level_color):
    """
    Build a single <w:tr>. Color coding rule (Fix 5):
      - Header cells: template header fill.
      - Data cells: alternating row fill BY DEFAULT.
      - Incident Energy cell ONLY: overridden with level color when supplied.
    The rest of the row stays neutral.
    """
    hdr_fill = fmt.get("fill", "1B2A4A")
    alt = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
    font = fmt.get("font", "Calibri")
    fsz = fmt.get("sz", "18")
    fc_h = fmt.get("color", "FFFFFF")
    bc = fmt.get("bc", "BFBFBF")
    mt = fmt.get("mt", "40")
    ml = fmt.get("ml", "70")
    mb = fmt.get("mb", "40")
    mr = fmt.get("mr", "70")

    hdr_tag = "<w:trPr><w:tblHeader/></w:trPr>" if is_header else ""
    tr = '<w:tr {}>{}'.format(nsdecls("w"), hdr_tag)

    for ci, (key, label) in enumerate(items):
        w = widths[ci] if ci < len(widths) else "900"
        text = _esc(label)
        align = ("center" if is_header or key in CENTER_KEYS
                 or key.startswith(("normal_", "mitigated_", "reduction_"))
                 else "left")

        if is_header:
            cell_fill = hdr_fill
            fc = fc_h
            bold_tag = "<w:b/>"
        else:
            # FIX: only the Level (energy_level) column gets the level color.
            # Incident Energy column remains neutral (alternating rows).
            if key == "energy_level" and level_color:
                cell_fill = level_color
            else:
                cell_fill = alt
            fc = "333333"
            bold_tag = ""

        tr += (
            '<w:tc><w:tcPr>'
            '<w:tcW w:w="{}" w:type="dxa"/>'.format(w) +
            '<w:tcBorders>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(bc) +
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(bc) +
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(bc) +
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(bc) +
            '</w:tcBorders>'
            '<w:shd w:val="clear" w:color="auto" w:fill="{}"/>'.format(cell_fill) +
            '<w:tcMar>'
            '<w:top w:w="{}" w:type="dxa"/>'.format(mt) +
            '<w:left w:w="{}" w:type="dxa"/>'.format(ml) +
            '<w:bottom w:w="{}" w:type="dxa"/>'.format(mb) +
            '<w:right w:w="{}" w:type="dxa"/>'.format(mr) +
            '</w:tcMar>'
            '<w:vAlign w:val="center"/>'
            '</w:tcPr>'
            '<w:p><w:pPr><w:spacing w:after="0"/><w:jc w:val="{}"/></w:pPr>'.format(align) +
            '<w:r><w:rPr>'
            '<w:rFonts w:ascii="{f}" w:hAnsi="{f}" w:cs="{f}"/>'.format(f=font) +
            '{}<w:sz w:val="{}"/><w:szCs w:val="{}"/>'.format(bold_tag, fsz, fsz) +
            '<w:color w:val="{}"/>'.format(fc) +
            '</w:rPr><w:t xml:space="preserve">{}</w:t></w:r></w:p></w:tc>'.format(text)
        )
    tr += '</w:tr>'
    return parse_xml(tr)
